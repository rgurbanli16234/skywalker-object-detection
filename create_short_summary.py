
from docx import Document
from docx.shared import Inches

doc = Document()

# Title
doc.add_heading('Skywalker Object Detection — Professor Submission', 0)

# Add sections
doc.add_heading('Model Name', level=1)
doc.add_paragraph('YOLOv8m (Medium)')

doc.add_heading('Final FPS', level=1)
doc.add_paragraph('30.6 FPS (GPU RTX 3050 Laptop)')

doc.add_heading('Final Accuracy Metrics', level=1)
metrics_table = doc.add_table(rows=1, cols=5)
hdr_cells = metrics_table.rows[0].cells
hdr_cells[0].text = 'Precision'
hdr_cells[1].text = 'Recall'
hdr_cells[2].text = 'F1-score'
hdr_cells[3].text = 'mAP50'
hdr_cells[4].text = 'mAP50-95'

row_cells = metrics_table.add_row().cells
row_cells[0].text = '0.00333'
row_cells[1].text = '1.0'
row_cells[2].text = '0.00664'
row_cells[3].text = '0.00516'
row_cells[4].text = '0.00309'

doc.add_heading('Why This Model Was Chosen', level=1)
doc.add_paragraph(
    'YOLOv8m was selected as it offers an excellent balance between accuracy and speed '
    'while being compatible with the available GPU resources (RTX 3050 Laptop with 3.68 GB VRAM). '
    'YOLOv8m provides robust performance for custom object detection tasks, featuring advanced '
    'architectural improvements over previous YOLO versions, including better feature extraction '
    'and multi-scale training capabilities.'
)

doc.save('submission/short_summary.docx')
